"""Export patent Markdown drafts to Word documents."""

from __future__ import annotations

import re
from pathlib import Path
from tempfile import TemporaryDirectory

from formula_utils import (
    append_formula_fallback,
    append_latex_omml,
    normalize_formula_markdown,
)
from tool_registry import register_tool


@register_tool(
    "export_word_document",
    "把最终 Markdown 转为 Word，并将 LaTeX 公式写成原生 Word 数学对象。",
    "Document production",
)
def export_markdown_to_docx(markdown: str, output_path: Path) -> Path:
    """Convert a Markdown-ish patent draft into a readable Word document."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:  # pragma: no cover - depends on local env
        raise RuntimeError("缺少 python-docx，请先运行 pip install -r requirements.txt") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_document(doc)

    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        for block in _iter_blocks(normalize_formula_markdown(markdown)):
            kind = block[0]
            if kind == "heading":
                _, level, text = block
                if "九、引用说明" in str(text):
                    doc.add_page_break()
                paragraph = doc.add_heading(_clean_inline(text), level=min(int(level), 3))
                _style_heading(paragraph)
                continue

            if kind == "list":
                _, items = block
                for item in items:
                    paragraph = doc.add_paragraph(style="List Bullet")
                    paragraph.add_run(_clean_inline(item))
                    _style_paragraph(paragraph)
                continue

            if kind == "mermaid":
                _, content = block
                image_path = _draw_mermaid_as_image(content, temp_path)
                if image_path:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(image_path), width=Cm(14.5))
                else:
                    _add_plain_paragraph(doc, "附图内容待补充。")
                continue

            if kind == "paragraph":
                _, text = block
                _add_rich_paragraph(doc, text)
                continue

            if kind == "formula":
                _, latex = block
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not append_latex_omml(paragraph, latex):
                    append_formula_fallback(paragraph, latex)
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if paragraph.style.name.startswith("Heading"):
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.save(output_path)
    return output_path


def _configure_document(doc: object) -> None:
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)


def _iter_blocks(markdown: str):
    lines = str(markdown or "").replace("\\n", "\n").splitlines()
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    list_items: list[str] = []
    paragraph_lines: list[str] = []
    in_formula = False
    formula_lines: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(line.strip() for line in paragraph_lines if line.strip())
            paragraph_lines = []
            if text:
                return ("paragraph", text)
        return None

    def flush_list():
        nonlocal list_items
        if list_items:
            items = list_items
            list_items = []
            return ("list", items)
        return None

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped == "$$":
            block = flush_paragraph()
            if block:
                yield block
            block = flush_list()
            if block:
                yield block
            if in_formula:
                in_formula = False
                formula = "\n".join(formula_lines).strip()
                formula_lines = []
                if formula:
                    yield ("formula", formula)
            else:
                in_formula = True
            continue

        if in_formula:
            formula_lines.append(line)
            continue

        single_formula = re.match(r"^\$\$(.+)\$\$$", stripped)
        if single_formula:
            block = flush_paragraph()
            if block:
                yield block
            block = flush_list()
            if block:
                yield block
            yield ("formula", single_formula.group(1).strip())
            continue

        fence = re.match(r"^```([a-zA-Z0-9_-]*)", line.strip())
        if fence:
            if in_code:
                in_code = False
                content = "\n".join(code_lines).strip()
                code_lines = []
                if code_lang.lower() == "mermaid" and content:
                    yield ("mermaid", content)
                elif content:
                    yield ("paragraph", content)
                code_lang = ""
            else:
                block = flush_paragraph()
                if block:
                    yield block
                block = flush_list()
                if block:
                    yield block
                in_code = True
                code_lang = fence.group(1)
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            block = flush_paragraph()
            if block:
                yield block
            block = flush_list()
            if block:
                yield block
            continue

        heading = re.match(r"^(#{1,6})\s*(.+)$", line.strip())
        if heading:
            block = flush_paragraph()
            if block:
                yield block
            block = flush_list()
            if block:
                yield block
            yield ("heading", min(len(heading.group(1)), 3), heading.group(2))
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line.strip()) or re.match(r"^\d+[.、]\s+(.+)$", line.strip())
        if bullet:
            block = flush_paragraph()
            if block:
                yield block
            list_items.append(bullet.group(1))
            continue

        block = flush_list()
        if block:
            yield block
        paragraph_lines.append(line)

    block = flush_paragraph()
    if block:
        yield block
    block = flush_list()
    if block:
        yield block
    if formula_lines:
        yield ("formula", "\n".join(formula_lines).strip())


def _add_plain_paragraph(doc: object, text: str) -> None:
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    _style_paragraph(paragraph)


def _add_rich_paragraph(doc: object, text: str) -> None:
    """Add a paragraph containing text and native inline Word equations."""
    cleaned = str(text or "")
    if not cleaned:
        return
    paragraph = doc.add_paragraph()
    position = 0
    for match in re.finditer(r"(?<!\\)\$([^$\n]+?)(?<!\\)\$", cleaned):
        prefix = _clean_inline(cleaned[position:match.start()])
        if prefix:
            paragraph.add_run(prefix)
        latex = match.group(1).strip()
        if not append_latex_omml(paragraph, latex):
            append_formula_fallback(paragraph, latex)
        position = match.end()
    suffix = _clean_inline(cleaned[position:])
    if suffix:
        paragraph.add_run(suffix)
    _style_paragraph(paragraph)


def _style_paragraph(paragraph: object) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)


def _style_heading(paragraph: object) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)


def _clean_inline(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"<!--.*?-->", "", value)
    value = value.replace("\\n", "\n")
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\*([^*]+)\*", r"\1", value)
    value = re.sub(r"^#+\s*", "", value.strip())
    value = value.replace("|", " | ")
    return re.sub(r"\s+", " ", value).strip()


def _draw_mermaid_as_image(content: str, temp_dir: Path) -> Path | None:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return None

    labels = _extract_mermaid_labels(content)
    if not labels:
        labels = ["数据获取", "处理分析", "方案生成", "结果输出"]

    width = 1400
    box_width = 940
    box_height = 86
    gap = 42
    top = 70
    height = top * 2 + len(labels) * box_height + max(0, len(labels) - 1) * gap
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = _load_font(size=28)
    small_font = _load_font(size=20)

    x = (width - box_width) // 2
    y = top
    for index, label in enumerate(labels, start=1):
        draw.rounded_rectangle(
            (x, y, x + box_width, y + box_height),
            radius=18,
            fill="#f3faf6",
            outline="#0f6b57",
            width=3,
        )
        prefix = f"S{index}"
        draw.text((x + 30, y + 26), prefix, fill="#0f6b57", font=small_font)
        draw.text((x + 105, y + 24), _shorten(label, 34), fill="#18201d", font=font)
        if index < len(labels):
            arrow_x = width // 2
            start_y = y + box_height + 8
            end_y = y + box_height + gap - 8
            draw.line((arrow_x, start_y, arrow_x, end_y), fill="#0f6b57", width=4)
            draw.polygon(
                [(arrow_x - 10, end_y - 2), (arrow_x + 10, end_y - 2), (arrow_x, end_y + 14)],
                fill="#0f6b57",
            )
        y += box_height + gap

    output = temp_dir / f"diagram_{abs(hash(content))}.png"
    image.save(output)
    return output


def _extract_mermaid_labels(content: str) -> list[str]:
    labels: list[str] = []
    for match in re.finditer(r"\[([^\[\]]{2,80})\]", content):
        label = _clean_inline(match.group(1))
        if label and label not in labels:
            labels.append(label)
    return labels[:10]


def _load_font(size: int):
    from PIL import ImageFont

    for path in (
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ):
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _shorten(text: str, limit: int) -> str:
    return text if len(text) <= limit else f"{text[:limit - 1]}…"
